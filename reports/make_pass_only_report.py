#!/usr/bin/env python3
"""Formal REAL PoC report — only PASS cases, polished charts, purpose/result/analysis."""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib import font_manager
from matplotlib.patches import FancyBboxPatch
import numpy as np
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
CHARTS = ROOT / "real_report" / "charts_v2"
DATA = ROOT / "real_report" / "data"
EV = Path("/tmp/l1-ev/evidence")
DOCX = Path("/Users/derry/Desktop/nlp/computing-power/L1_算力调度真实调度测试报告_20260811.docx")

# palette
C_NAVY = "#0B3D5C"
C_BLUE = "#1F6AA5"
C_TEAL = "#1A9B8E"
C_GREEN = "#2E8B57"
C_ORANGE = "#E07A3D"
C_RED = "#C0392B"
C_BG = "#F7FAFC"
C_GRID = "#E2E8F0"
C_TEXT = "#1A202C"

BLUE = "0B3D5C"
PALE = "F0F7FB"
GRAY = "64748B"
FONT = "Songti SC"

# Only passed cases (including expected-reject passes)
CASES = [
    {
        "id": "DT01",
        "title": "两地共同执行的分片批量推理",
        "level": "A",
        "purpose": "验证新加坡控制面可将冻结 ResNet-50 父任务拆成 8 个分片，在海南与重庆真实 GPU 上共同执行，并在控制面汇聚结果；要求连续 3 轮成功，样本无缺失/重复，且与冻结参考输出一致。",
        "result": "三轮父任务 POC-REAL-0004/0005/0006 均为 SUCCEEDED（8/8）。每轮海南 6 片、重庆 2 片；缺失 0、重复 0。复核轮 POC-REAL-0001 参考比对 compared=4096、mismatch=0、pass=True。判定 PASS-REAL。",
        "analysis": "调度呈现“海南优先、重庆补位”：前两片落海南后派重庆，之后继续回填海南。说明硬约束与策略在真实多节点下可闭环。参考比对全过，证明不是“只跑通接口”，而是业务输出可校验。海南现场仅 1×4090，故海南分片串行偏多；不影响“两地均参与”结论。",
        "chart": "dt01_combo.png",
        "metrics": [
            ["轮次", "父任务", "成功分片", "海南", "重庆", "结果"],
            ["1", "POC-REAL-0004", "8/8", "6", "2", "SUCCEEDED"],
            ["2", "POC-REAL-0005", "8/8", "6", "2", "SUCCEEDED"],
            ["3", "POC-REAL-0006", "8/8", "6", "2", "SUCCEEDED"],
            ["复核", "POC-REAL-0001", "8/8", "—", "—", "SUCCEEDED + 参考比对通过"],
        ],
    },
    {
        "id": "DT03",
        "title": "单卡 16GB 显存硬约束",
        "level": "A",
        "purpose": "验证调度器按“单卡可用显存”而非“整机显存池”过滤候选：16GB 任务不得派往重庆 12GB 的 4070，只能选择海南 24GB 的 4090。",
        "result": "父任务 POC-REAL-0002：1/1 分片 SUCCEEDED，执行节点仅为海南-GPU1。重庆四张 12GB 卡均未入选。判定 PASS-REAL。",
        "analysis": "结果符合 L1 边界：不允许把 4×12GB 拼成 48GB。过滤发生在候选阶段，避免 OOM 后再回退。与文档 DT03/TC10 口径一致。",
        "chart": "dt03_filter.png",
        "metrics": [
            ["候选 GPU", "显存(GB)", "≥16GB?", "调度结果"],
            ["重庆-GPU1~4", "12", "否", "排除"],
            ["海南-GPU1", "24", "是", "选中并成功"],
        ],
    },
    {
        "id": "TC01",
        "title": "正常跨境接入与完整性校验",
        "level": "A",
        "purpose": "验证任务包在创建可执行任务前完成完整性校验（manifest/数据 SHA），校验通过后可正常登记并调度到真实节点。",
        "result": "父任务 POC-REAL-0003：integrity ok；4/4 分片成功（海南 S01–S02，重庆 S03–S04）。判定 PASS-REAL。",
        "analysis": "接入链路“校验→登记→分片下发→回传”完整。与负向 TC04 对照：哈希正确才放行，错误则拒绝，形成正反闭环。",
        "chart": "tc01_flow.png",
        "metrics": [
            ["子任务", "区域", "状态", "样本数"],
            ["S01", "海南", "SUCCEEDED", "512"],
            ["S02", "海南", "SUCCEEDED", "512"],
            ["S03", "重庆", "SUCCEEDED", "512"],
            ["S04", "重庆", "SUCCEEDED", "512"],
        ],
    },
    {
        "id": "TC02",
        "title": "身份认证失败",
        "level": "A",
        "purpose": "验证使用错误 Bearer 访问节点 agent 时被拒绝，控制面记录审计，且不创建可执行 GPU 任务。",
        "result": "海南与重庆 agent 对 invalid token 均返回 HTTP 401；父任务状态 REJECTED；未启动 ResNet。判定 PASS-REAL（预期拒绝）。",
        "analysis": "安全负向用例以“拒绝成功”为通过标准。双节点一致 401，说明鉴权不依赖单点。审计含 AUTH_FAIL，满足可追溯要求。",
        "chart": "security_panel.png",
        "metrics": [
            ["检查项", "观测"],
            ["海南 agent", "HTTP 401"],
            ["重庆 agent", "HTTP 401"],
            ["GPU 任务", "未创建"],
            ["审计事件", "AUTH_FAIL / COMPLETE(REJECTED)"],
        ],
    },
    {
        "id": "TC03",
        "title": "非授权端口或路径",
        "level": "A",
        "purpose": "验证对非授权端口的探针无法建立业务连接，系统记录阻断证据且不创建执行任务。",
        "result": "探针 127.0.0.1:59999 Connection refused，blocked=true；任务 REJECTED。判定 PASS-REAL（预期拒绝）。",
        "analysis": "在无正式 VPN/防火墙编排条件下，以显式非授权端口探针验证“默认拒绝”。与 TC02 共同覆盖访问控制边界。",
        "chart": "security_panel.png",
        "metrics": [
            ["检查项", "观测"],
            ["探针目标", "127.0.0.1:59999"],
            ["结果", "Connection refused / blocked=true"],
            ["GPU 任务", "未创建"],
            ["审计事件", "PROBE_BLOCKED"],
        ],
    },
    {
        "id": "TC04",
        "title": "数据完整性异常",
        "level": "A",
        "purpose": "验证当期望数据集哈希被故意篡改时，系统拒绝创建可执行任务并保留 HASH_MISMATCH 审计。",
        "result": "故意错误期望哈希触发拒绝；HTTP 业务响应成功返回拒绝结论；未下发 GPU。判定 PASS-REAL（预期拒绝）。",
        "analysis": "与 TC01 形成对照：正确哈希放行、错误哈希拒绝。防止“脏数据入库后才失败”。",
        "chart": "security_panel.png",
        "metrics": [
            ["检查项", "观测"],
            ["触发方式", "expect_tar_hash 置为全 0"],
            ["系统行为", "拒绝建任务"],
            ["审计事件", "HASH_MISMATCH"],
            ["GPU 任务", "未创建"],
        ],
    },
    {
        "id": "TC10",
        "title": "单卡显存硬约束",
        "level": "A",
        "purpose": "验证资源声明中的单卡显存下限在真实调度中生效（与 DT03 同源约束）。",
        "result": "由真实 DT03 覆盖：16GB 任务仅落海南。判定 PASS-REAL（mapped:DT03）。",
        "analysis": "映射覆盖符合方案“核心场景可映射到原子用例”的写法；真实证据与 DT03 同一父任务链。",
        "chart": "dt03_filter.png",
        "metrics": [
            ["映射来源", "真实父任务", "结论"],
            ["mapped:DT03", "POC-REAL-0002", "PASS-REAL"],
        ],
    },
    {
        "id": "TC13",
        "title": "海南本地推理",
        "level": "A",
        "purpose": "验证海南节点可真实执行 ResNet 分片推理并回传 predictions。",
        "result": "DT01 各轮均有海南分片 SUCCEEDED（如 S01/S02/S05–S08）。判定 PASS-REAL（mapped:DT01）。",
        "analysis": "海南 1×4090 承担主要分片，证明本地推理链路（agent→torch/CUDA→回传）可用。",
        "chart": "dt01_combo.png",
        "metrics": [
            ["覆盖来源", "典型分片", "状态"],
            ["mapped:DT01", "S01/S02/S05–S08", "SUCCEEDED"],
        ],
    },
    {
        "id": "TC14",
        "title": "重庆异地推理",
        "level": "A",
        "purpose": "验证重庆节点可作为异地执行点完成分片推理并回传。",
        "result": "DT01 各轮 S03/S04 派重庆并 SUCCEEDED。判定 PASS-REAL（mapped:DT01）。",
        "analysis": "异地执行已打通；链路为公网可达而非正式 OTN SLA，本报告只确认功能成功，不宣称专线指标达标。",
        "chart": "dt01_combo.png",
        "metrics": [
            ["覆盖来源", "典型分片", "状态"],
            ["mapped:DT01", "S03/S04", "SUCCEEDED"],
        ],
    },
    {
        "id": "TC15",
        "title": "批量推理与队列",
        "level": "A",
        "purpose": "验证批量分片任务可连续排队/下发并全部完成，形成可追踪闭环。",
        "result": "DT01 连续 3 轮、每轮 8 分片全部成功。判定 PASS-REAL（mapped:DT01）。",
        "analysis": "批量场景稳定性体现在“连续三轮零失败”。分片串行下发适合当前单卡海南拓扑。",
        "chart": "dt01_combo.png",
        "metrics": [
            ["轮次", "分片", "失败", "结果"],
            ["3", "8/轮", "0", "全部 SUCCEEDED"],
        ],
    },
    {
        "id": "TC19",
        "title": "结果正确性",
        "level": "A",
        "purpose": "验证汇聚输出与冻结参考结果一致，避免“任务成功但结果不可信”。",
        "result": "reference_match：compared=4096，mismatch=0，pass=True。判定 PASS-REAL（mapped:DT01）。",
        "analysis": "top1_class 全量一致是本期最强正确性证据。文件级 SHA 可能因 shard 元数据不同而不同，业务以类别一致为准。",
        "chart": "tc19_match.png",
        "metrics": [
            ["指标", "值"],
            ["compared", "4096"],
            ["mismatch", "0"],
            ["pass", "True"],
        ],
    },
]


def setup_font():
    for c in ["/System/Library/Fonts/Supplemental/Songti.ttc", "/System/Library/Fonts/STHeiti Light.ttc"]:
        if Path(c).exists():
            font_manager.fontManager.addfont(c)
            plt.rcParams["font.family"] = font_manager.FontProperties(fname=c).get_name()
            break
    plt.rcParams.update({
        "axes.unicode_minus": False,
        "figure.facecolor": "white",
        "axes.facecolor": C_BG,
        "axes.edgecolor": C_GRID,
        "axes.labelcolor": C_TEXT,
        "text.color": C_TEXT,
        "xtick.color": C_TEXT,
        "ytick.color": C_TEXT,
        "grid.color": C_GRID,
        "grid.linewidth": 0.8,
        "axes.titlesize": 13,
        "axes.labelsize": 10,
    })


def style_ax(ax):
    ax.grid(axis="y", alpha=0.7)
    ax.set_axisbelow(True)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(C_GRID)
    ax.spines["bottom"].set_color(C_GRID)


def make_charts():
    CHARTS.mkdir(parents=True, exist_ok=True)
    setup_font()
    dt01 = json.loads((DATA / "dt01_rounds.json").read_text(encoding="utf-8"))

    # cover summary
    fig = plt.figure(figsize=(11.2, 4.2))
    gs = fig.add_gridspec(1, 2, width_ratios=[1.15, 1], wspace=0.28)
    ax0 = fig.add_subplot(gs[0])
    ax1 = fig.add_subplot(gs[1])
    labels = ["正向通过", "负向通过\n(预期拒绝)", "映射覆盖"]
    vals = [3, 3, 5]  # DT01/DT03/TC01 ; TC02/03/04 ; TC10/13/14/15/19
    colors = [C_TEAL, C_ORANGE, C_BLUE]
    bars = ax0.bar(labels, vals, color=colors, width=0.55, zorder=3)
    for b, v in zip(bars, vals):
        ax0.text(b.get_x() + b.get_width() / 2, v + 0.12, str(v), ha="center", fontweight="bold")
    ax0.set_ylim(0, 6)
    ax0.set_ylabel("用例数")
    ax0.set_title("本期写入报告的通过用例构成（共 11 项）")
    style_ax(ax0)

    # donut
    sizes = [11]
    wedges, _ = ax1.pie([11], colors=[C_TEAL], startangle=90, wedgeprops=dict(width=0.42, edgecolor="white"))
    ax1.text(0, 0.08, "11", ha="center", va="center", fontsize=28, fontweight="bold", color=C_NAVY)
    ax1.text(0, -0.28, "全部通过", ha="center", va="center", fontsize=11, color=C_TEXT)
    ax1.set_title("报告收录范围")
    fig.suptitle("跨境算力调度 PoC · 真实调度通过用例总览", fontsize=14, fontweight="bold", color=C_NAVY, y=1.02)
    fig.tight_layout()
    fig.savefig(CHARTS / "summary_overview.png", dpi=180, bbox_inches="tight")
    plt.close()

    # architecture prettier
    fig, ax = plt.subplots(figsize=(11.2, 3.8))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 4)
    ax.axis("off")
    boxes = [
        (0.4, 1.2, 2.6, 1.6, "新加坡控制面\nUI :8080\n调度 / 汇聚 / 审计", C_BLUE),
        (4.0, 1.2, 2.6, 1.6, "调度决策\n硬约束过滤\n海南优先策略", "#C05621"),
        (7.6, 2.15, 3.6, 1.25, "海南 agent :8000\n1×RTX 4090 · ResNet", C_TEAL),
        (7.6, 0.45, 3.6, 1.25, "重庆 agent :8000\n4×RTX 4070 · ResNet", C_GREEN),
    ]
    for x, y, w, h, text, color in boxes:
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.04,rounding_size=0.15",
                                    facecolor=color, edgecolor="white", linewidth=2, alpha=0.92))
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", color="white", fontsize=10, fontweight="bold")
    ax.annotate("", xy=(4.0, 2.0), xytext=(3.0, 2.0), arrowprops=dict(arrowstyle="-|>", color=C_NAVY, lw=2))
    ax.annotate("", xy=(7.6, 2.7), xytext=(6.6, 2.2), arrowprops=dict(arrowstyle="-|>", color=C_NAVY, lw=2))
    ax.annotate("", xy=(7.6, 1.0), xytext=(6.6, 1.7), arrowprops=dict(arrowstyle="-|>", color=C_NAVY, lw=2))
    ax.set_title("真实调度数据面（本期实测拓扑）", fontsize=13, fontweight="bold", color=C_NAVY, pad=8)
    fig.tight_layout()
    fig.savefig(CHARTS / "architecture.png", dpi=180, bbox_inches="tight")
    plt.close()

    # DT01 combo: timeline + region bars
    fig = plt.figure(figsize=(11.2, 6.6))
    gs = fig.add_gridspec(2, 1, height_ratios=[1.35, 1], hspace=0.35)
    ax = fig.add_subplot(gs[0])
    r1 = dt01["R1"]["shards"]
    t0 = datetime.fromisoformat(r1[0][2])
    cmap = {"海南": C_BLUE, "重庆": C_TEAL}
    for i, (sid, region, st, ed) in enumerate(r1):
        s = (datetime.fromisoformat(st) - t0).total_seconds()
        e = (datetime.fromisoformat(ed) - t0).total_seconds()
        ax.barh(i, max(e - s, 0.8), left=s, height=0.62, color=cmap[region], alpha=0.9, zorder=3)
        ax.text(e + 0.5, i, f"{region} · {e-s:.0f}s", va="center", fontsize=8.5, color=C_TEXT)
    ax.set_yticks(range(len(r1)))
    ax.set_yticklabels([x[0] for x in r1])
    ax.invert_yaxis()
    ax.set_xlabel("相对时间（秒）")
    ax.set_title("DT01 第1轮真实分片时间线（POC-REAL-0004）")
    style_ax(ax)
    ax.legend(handles=[mpatches.Patch(color=c, label=k) for k, c in cmap.items()], loc="lower right", frameon=False)

    ax2 = fig.add_subplot(gs[1])
    rounds = ["第1轮", "第2轮", "第3轮"]
    hn, cq = [], []
    for key in ("R1", "R2", "R3"):
        from collections import Counter
        regs = Counter(s[1] for s in dt01[key]["shards"])
        hn.append(regs["海南"]); cq.append(regs["重庆"])
    x = np.arange(len(rounds))
    w = 0.34
    b1 = ax2.bar(x - w / 2, hn, w, color=C_BLUE, label="海南", zorder=3)
    b2 = ax2.bar(x + w / 2, cq, w, color=C_TEAL, label="重庆", zorder=3)
    for bars in (b1, b2):
        for b in bars:
            ax2.text(b.get_x() + b.get_width() / 2, b.get_height() + 0.1, f"{int(b.get_height())}", ha="center", fontsize=9)
    ax2.set_xticks(x); ax2.set_xticklabels(rounds)
    ax2.set_ylim(0, 9); ax2.set_ylabel("成功分片数")
    ax2.set_title("连续三轮两地参与（每轮均为 8/8 成功）")
    ax2.legend(frameon=False)
    style_ax(ax2)
    fig.suptitle("DT01 主流程可视化", fontsize=14, fontweight="bold", color=C_NAVY)
    fig.tight_layout()
    fig.savefig(CHARTS / "dt01_combo.png", dpi=180, bbox_inches="tight")
    plt.close()

    # DT03 filter
    fig, ax = plt.subplots(figsize=(10.5, 4.2))
    nodes = ["重庆\nGPU1", "重庆\nGPU2", "重庆\nGPU3", "重庆\nGPU4", "海南\nGPU1"]
    mem = [12, 12, 12, 12, 24]
    cols = [C_RED, C_RED, C_RED, C_RED, C_TEAL]
    bars = ax.bar(nodes, mem, color=cols, width=0.58, zorder=3, alpha=0.92)
    ax.axhline(16, color=C_ORANGE, ls="--", lw=2, label="需求阈值 16GB")
    for b, v, ok in zip(bars, mem, [False] * 4 + [True]):
        ax.text(b.get_x() + b.get_width() / 2, v + 0.5, f"{v}GB\n{'入选' if ok else '排除'}",
                ha="center", fontsize=8.5, color=C_TEAL if ok else C_RED)
    ax.set_ylim(0, 30); ax.set_ylabel("单卡显存 (GB)")
    ax.set_title("DT03 / TC10：16GB 硬约束候选过滤（真实资源）")
    ax.legend(frameon=False, loc="upper left")
    style_ax(ax)
    fig.tight_layout()
    fig.savefig(CHARTS / "dt03_filter.png", dpi=180, bbox_inches="tight")
    plt.close()

    # TC01 flow
    fig, axes = plt.subplots(1, 2, figsize=(10.8, 4.0))
    axes[0].set_xlim(0, 10); axes[0].set_ylim(0, 4); axes[0].axis("off")
    steps = [(0.3, "完整性\n校验"), (2.7, "任务\n登记"), (5.1, "分片\n下发"), (7.5, "结果\n回传")]
    for i, (x, t) in enumerate(steps):
        axes[0].add_patch(FancyBboxPatch((x, 1.2), 1.9, 1.5, boxstyle="round,pad=0.03,rounding_size=0.12",
                                         facecolor=C_TEAL if i else C_BLUE, edgecolor="white", lw=2))
        axes[0].text(x + 0.95, 1.95, t, ha="center", va="center", color="white", fontsize=10, fontweight="bold")
        if i < 3:
            axes[0].annotate("", xy=(x + 2.25, 1.95), xytext=(x + 2.05, 1.95),
                             arrowprops=dict(arrowstyle="-|>", color=C_NAVY, lw=1.8))
    axes[0].set_title("TC01 正向接入链路")
    axes[1].bar(["海南\nS01-S02", "重庆\nS03-S04"], [2, 2], color=[C_BLUE, C_TEAL], width=0.5, zorder=3)
    axes[1].set_ylim(0, 3.2); axes[1].set_ylabel("成功分片")
    axes[1].set_title("TC01 分片落点（4/4）")
    for i, v in enumerate([2, 2]):
        axes[1].text(i, v + 0.08, str(v), ha="center", fontweight="bold")
    style_ax(axes[1])
    fig.suptitle("TC01 正常跨境接入", fontsize=13, fontweight="bold", color=C_NAVY)
    fig.tight_layout()
    fig.savefig(CHARTS / "tc01_flow.png", dpi=180, bbox_inches="tight")
    plt.close()

    # security panel
    fig, ax = plt.subplots(figsize=(10.5, 4.2))
    xs = np.arange(3)
    ax.bar(xs, [1, 1, 1], color=[C_RED, C_ORANGE, C_RED], width=0.5, zorder=3, alpha=0.9)
    ax.set_xticks(xs)
    ax.set_xticklabels(["TC02\n错误 Bearer → 401", "TC03\n非授权端口阻断", "TC04\n哈希不一致拒绝"])
    ax.set_ylim(0, 1.45)
    ax.set_yticks([0, 1]); ax.set_yticklabels(["放行", "拒绝"])
    for i, note in enumerate(["双 agent 401\n未建 GPU 任务", "59999 refused\n未建 GPU 任务", "HASH_MISMATCH\n未建 GPU 任务"]):
        ax.text(i, 1.08, note, ha="center", fontsize=8.5, color=C_TEXT)
    ax.set_title("安全负向用例：预期拒绝即通过")
    style_ax(ax)
    fig.tight_layout()
    fig.savefig(CHARTS / "security_panel.png", dpi=180, bbox_inches="tight")
    plt.close()

    # TC19 match
    fig, axes = plt.subplots(1, 2, figsize=(10.5, 4.0))
    axes[0].bar(["比对样本", "不一致"], [4096, 0], color=[C_TEAL, C_RED], width=0.45, zorder=3)
    axes[0].set_ylim(0, 4600)
    axes[0].text(0, 4200, "4096", ha="center", fontweight="bold", color=C_TEAL)
    axes[0].text(1, 180, "0", ha="center", fontweight="bold", color=C_RED)
    axes[0].set_title("参考输出 top1 比对")
    style_ax(axes[0])
    # class dist top
    from collections import Counter
    c = Counter()
    pred = EV / "POC-REAL-0001" / "merged_predictions.jsonl"
    if pred.exists():
        for line in open(pred):
            o = json.loads(line)
            c[o["top1_class"]] += 1
    top = c.most_common(6) if c else [(0, 0)]
    axes[1].bar([str(k) for k, _ in top], [v for _, v in top], color=C_BLUE, alpha=0.9, zorder=3)
    axes[1].set_title("汇聚结果 top1_class 分布（节选）")
    axes[1].set_xlabel("class id"); axes[1].set_ylabel("样本数")
    style_ax(axes[1])
    fig.suptitle("TC19 结果正确性", fontsize=13, fontweight="bold", color=C_NAVY)
    fig.tight_layout()
    fig.savefig(CHARTS / "tc19_match.png", dpi=180, bbox_inches="tight")
    plt.close()

    print("charts ->", CHARTS)


# -------- docx helpers --------

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd"); tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color="000000", size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_margins(cell, t=70, s=70, b=70, e=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar"); tc_pr.append(tc_mar)
    for m, v in (("top", t), ("start", s), ("bottom", b), ("end", e)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}"); tc_mar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def borders(table, color="D0E3F0", size="6"):
    tbl_pr = table._tbl.tblPr
    b = tbl_pr.first_child_found_in("w:tblBorders")
    if b is None:
        b = OxmlElement("w:tblBorders"); tbl_pr.append(b)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = b.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}"); b.append(el)
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), color)


def add_table(doc, headers, rows, widths=None, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(table)
    hdr = table.rows[0]
    tr_pr = hdr._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true"); tr_pr.append(th)
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, bold=True, color="FFFFFF", size=font_size)
        shade(hdr.cells[i], BLUE); set_margins(hdr.cells[i])
        if widths: hdr.cells[i].width = Cm(widths[i])
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, v in enumerate(row):
            set_cell_text(cells[i], v, size=font_size)
            if ridx % 2 == 0: shade(cells[i], PALE)
            set_margins(cells[i])
            if widths: cells[i].width = Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(); p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = FONT; run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.color.rgb = RGBColor.from_string(BLUE if level == 1 else "1F6AA5")


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = FONT; run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def add_labeled(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    r1 = p.add_run(label)
    r1.bold = True; r1.font.color.rgb = RGBColor.from_string(BLUE)
    r1.font.name = FONT; r1._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r2 = p.add_run(text)
    r2.font.name = FONT; r2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def add_callout(doc, title, body, fill="E6F4F1"):
    t = doc.add_table(rows=1, cols=1)
    borders(t, "9ED5CB", "10")
    cell = t.cell(0, 0); shade(cell, fill); set_margins(cell, 130, 150, 130, 150)
    p = cell.paragraphs[0]
    r = p.add_run(title); r.bold = True; r.font.color.rgb = RGBColor.from_string(BLUE)
    r.font.name = FONT; r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    p2 = cell.add_paragraph(body)
    for run in p2.runs:
        run.font.name = FONT; run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc, path, caption, width_cm=16.0):
    if not Path(path).exists():
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inline = p.add_run().add_picture(str(path), width=Cm(width_cm))
    inline._inline.docPr.set("descr", caption)
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(caption); r.italic = True; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(GRAY)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    a = OxmlElement("w:fldChar"); a.set(qn("w:fldCharType"), "begin")
    b = OxmlElement("w:instrText"); b.set(qn("xml:space"), "preserve"); b.text = "PAGE"
    c = OxmlElement("w:fldChar"); c.set(qn("w:fldCharType"), "end")
    run._r.append(a); run._r.append(b); run._r.append(c)
    paragraph.add_run(" 页")


def build_doc():
    make_charts()
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.9); sec.bottom_margin = Cm(1.7)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Normal"].font.size = Pt(10.5)
    for name, size in [("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 11)]:
        st = styles[name]
        st.font.name = FONT; st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(size); st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(BLUE)
    footer = sec.footer.paragraphs[0]
    footer.text = "L1 真实调度测试报告（仅通过用例） | 2026-08-11 | "
    footer.runs[0].font.size = Pt(8); footer.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    add_page_number(sec.footer.add_paragraph())

    # cover
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(64)
    r = p.add_run("跨境算力调度 PoC"); r.font.size = Pt(16); r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("L1 算力调度真实调度测试报告"); r.font.size = Pt(24); r.bold = True
    r.font.color.rgb = RGBColor.from_string("0B3D5C")
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run("仅收录已通过用例 · 含目的 / 结果 / 分析与可视化")
    r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string(GRAY)

    add_table(doc, ["项目", "内容"], [
        ["报告编号", "L1-SCHED-REAL-20260811"],
        ["执行日期", "2026年8月11日"],
        ["测试性质", "真实调度（新加坡控制面 → 海南/重庆 GPU ResNet 推理）"],
        ["收录范围", "仅 PASS-REAL / 预期拒绝通过 / 映射覆盖通过，共 11 项"],
        ["未收录", "未执行、未取证、失败项一律不写入本报告"],
        ["资源边界", "海南 RTX 4090×1（实测）；重庆 RTX 4070×4"],
    ], widths=[3.2, 13.0], font_size=9)

    add_heading(doc, "执行摘要", 1)
    add_callout(doc, "结论：本期收录的 11 项用例全部通过",
                "主流程 DT01 连续三轮 8/8 成功且参考比对全过；DT03/TC10 显存硬约束正确；"
                "TC01 正向接入通过；TC02/TC03/TC04 负向安全场景按预期拒绝。"
                "TC13/14/15/19 由 DT01 映射覆盖。未通过或未做的用例不进入本报告。")
    add_figure(doc, CHARTS / "summary_overview.png", "图1  通过用例构成总览", 15.8)
    add_figure(doc, CHARTS / "architecture.png", "图2  真实调度架构", 15.8)
    add_table(doc, ["用例", "级别", "场景", "判定"], [
        [c["id"], c["level"], c["title"], "PASS-REAL" + ("（预期拒绝）" if "拒绝" in c["result"] else ("（映射）" if "mapped" in c["result"] else ""))]
        for c in CASES
    ], widths=[1.6, 1.2, 8.5, 4.5], font_size=8)

    add_heading(doc, "1. 测试环境与方法", 1)
    add_body(doc, "依据《跨境算力调度 PoC 测试实施方案 v0724》。主业务为冻结 ResNet-50 + 4096 样本 / 8 分片。"
             "新加坡控制面负责调度与汇聚，海南/重庆 node agent 在真实 GPU 执行推理并回传 predictions。")
    add_table(doc, ["组件", "部署", "说明"], [
        ["控制面+UI", "新加坡 43.106.50.98:8080", "选点、分片下发、汇聚、审计"],
        ["海南 agent", "隧道映射 :18000 → :8000", "1×4090；torch cu124"],
        ["重庆 agent", "218.201.8.129:8000", "4×4070"],
        ["冻结包", "/opt/l1-poc-bundle", "数据/权重/reference"],
    ], widths=[3, 5.5, 7.5], font_size=8.5)

    add_heading(doc, "2. 通过用例详述", 1)
    fig_no = 3
    for case in CASES:
        add_heading(doc, f"{case['id']}  {case['title']}", 2)
        add_labeled(doc, "【目的】", case["purpose"])
        add_labeled(doc, "【结果】", case["result"])
        add_labeled(doc, "【分析】", case["analysis"])
        headers, *rows = case["metrics"]
        add_table(doc, headers, rows, font_size=8)
        chart = CHARTS / case["chart"]
        if chart.exists():
            # avoid repeating same chart caption noise too much — still include
            add_figure(doc, chart, f"图{fig_no}  {case['id']} 可视化", 15.6)
            fig_no += 1
        # evidence snippet if available
        ev = EV / "case_evidence" / f"{case['id']}.json"
        if ev.exists() and case["id"] in ("DT01", "DT03", "TC01", "TC02", "TC03", "TC04"):
            d = json.load(open(ev))
            add_table(doc, ["证据字段", "值"], [
                ["parent_task_id", d.get("parent_task_id", "")],
                ["execution_time", d.get("execution_time", "")],
                ["execution_node", d.get("execution_node", "")],
                ["test_conclusion", d.get("test_conclusion", "")[:120]],
            ], widths=[4, 12], font_size=8)

    add_heading(doc, "3. 证据与复现", 1)
    add_table(doc, ["材料", "位置"], [
        ["本报告", DOCX.name],
        ["证据包", "l1-poc-evidence-latest.tgz"],
        ["图表", "extracted/evidence/real_report/charts_v2/"],
        ["源码", "extracted/webapp/real/"],
        ["现场证据", "/opt/l1-scheduler-ui/evidence/"],
    ], widths=[3.5, 12.5], font_size=8.5)

    add_heading(doc, "4. 签署", 1)
    add_table(doc, ["角色", "姓名", "日期", "签字"], [
        ["测试执行", "", "2026-08-11", ""],
        ["结果复核", "", "", ""],
        ["项目确认", "", "", ""],
    ], widths=[4, 4, 4, 4], font_size=9)

    doc.save(str(DOCX))
    print("saved", DOCX, DOCX.stat().st_size)


if __name__ == "__main__":
    # ensure evidence
    if not (EV / "case_evidence" / "DT01.json").exists():
        import tarfile
        tarfile.open("/Users/derry/Desktop/nlp/computing-power/l1-poc-evidence-latest.tgz").extractall("/tmp/l1-ev")
    build_doc()
