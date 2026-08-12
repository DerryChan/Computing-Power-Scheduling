#!/usr/bin/env python3
"""生成跨境算力调度算法测试报告（最终版）。"""

from __future__ import annotations

import json
import sys
from datetime import datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt  # noqa: E402
import numpy as np  # noqa: E402
from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Cm, Pt, RGBColor  # noqa: E402

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from scheduler.experiment import assert_experiment_health, run_paper_experiment  # noqa: E402

REPORTS = ROOT / "reports"
DATA = REPORTS / "data"
CHARTS = REPORTS / "charts"
C_NAVY = "1d4e89"
C_TEAL = "#0f766e"
METHOD = "本文方法（动态权重多目标调度）"


def setup_font() -> None:
    plt.rcParams["font.sans-serif"] = ["Arial Unicode MS", "PingFang SC", "Heiti SC", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False


def set_run_font(run, size=11, bold=False, color=None) -> None:
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    r = run._element.get_or_add_rPr()
    rFonts = r.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "宋体")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text, *, size=11, bold=False, center=False, space_after=8):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size=16 if level == 1 else 13, bold=True, color=C_NAVY)
    return h


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=9, bold=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9)
    doc.add_paragraph()
    return table


def add_figure(doc, path: Path, caption: str, width_cm=15.2):
    if not path.exists():
        add_para(doc, f"[缺图] {caption}", size=10)
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, caption, size=10, center=True, space_after=12)


def short_name(name: str) -> str:
    return name.replace("本文方法（动态权重多目标调度）", "动态权重多目标").replace("本文方法（全模块）", "全模块")


def chart_bars(summaries: dict, out: Path) -> None:
    names = list(summaries.keys())
    short = [short_name(n) for n in names]
    fig, axes = plt.subplots(1, 3, figsize=(12.5, 4.2))
    colors = ["#94a3b8", "#64748b", "#1d4e89", "#b45309", "#7c2d12", "#0f766e"]
    series = [
        ([summaries[n]["success_rate_pct"] for n in names], "调度成功率 (%)", "%"),
        ([summaries[n]["avg_latency_ms"] for n in names], "平均时延 (ms)", "ms"),
        ([summaries[n]["avg_cost"] for n in names], "平均成本 (元)", "元"),
    ]
    for ax, (vals, title, ylabel) in zip(axes, series):
        bars = ax.bar(range(len(vals)), vals, color=colors[: len(vals)])
        ax.set_xticks(range(len(vals)))
        ax.set_xticklabels(short, fontsize=7, rotation=18, ha="right")
        ax.set_title(title)
        ax.set_ylabel(ylabel)
        for b, v in zip(bars, vals):
            ax.text(b.get_x() + b.get_width() / 2, v, f"{v:.1f}", ha="center", va="bottom", fontsize=8)
    fig.suptitle("算法对比关键指标")
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)


def chart_radar(summaries: dict, out: Path) -> None:
    labels = ["成功率", "峰值GPU利用率", "时延满意度", "成本经济度", "绿电比例"]
    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist() + [0]
    fig, ax = plt.subplots(figsize=(7.2, 6.2), subplot_kw=dict(polar=True))
    colors = ["#64748b", "#94a3b8", "#1d4e89", "#b45309", "#7c2d12", "#0f766e"]
    for i, (name, s) in enumerate(summaries.items()):
        lat_sat = max(0, min(100, 140 - float(s["avg_latency_ms"])))
        cost_eco = max(0, min(100, 100 - float(s["avg_cost"])))
        vals = [
            float(s["success_rate_pct"]),
            float(s["gpu_util_pct"]),
            lat_sat,
            cost_eco,
            float(s["avg_green_pct"]),
        ] + [float(s["success_rate_pct"])]
        ax.plot(
            angles,
            vals,
            color=colors[i % len(colors)],
            linewidth=2.4 if name == METHOD else 1.3,
            label=short_name(name),
        )
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels)
    ax.set_ylim(0, 100)
    ax.legend(loc="upper right", bbox_to_anchor=(1.45, 1.12), fontsize=8)
    ax.set_title("多维度对比")
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)


def chart_dist(rows: list, out: Path) -> None:
    from collections import Counter

    c = Counter(r["selected"] for r in rows if r.get("status") == "SCHEDULED" and r.get("selected"))
    names = ["重庆", "海南", "香港", "新加坡", "新疆"]
    vals = [c.get(n, 0) for n in names]
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.bar(names, vals, color=["#1d4e89", "#0f766e", "#b45309", "#0ea5e9", "#65a30d"])
    ax.set_ylabel("成功任务数")
    ax.set_title("动态权重多目标 · 任务落点分布")
    for i, v in enumerate(vals):
        ax.text(i, v + 0.05, str(v), ha="center")
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)


def chart_ablation(ablation: dict, out: Path) -> None:
    names = list(ablation.keys())
    short = [short_name(n) for n in names]
    fig, axes = plt.subplots(1, 3, figsize=(12, 4))
    for ax, key, title in [
        (axes[0], "success_rate_pct", "成功率 (%)"),
        (axes[1], "avg_latency_ms", "平均时延 (ms)"),
        (axes[2], "avg_cost", "平均成本 (元)"),
    ]:
        vals = [ablation[n][key] for n in names]
        ax.plot(range(len(vals)), vals, marker="o", color=C_TEAL)
        ax.set_xticks(range(len(vals)))
        ax.set_xticklabels(short, fontsize=7, rotation=20, ha="right")
        ax.set_title(title)
    fig.suptitle("消融实验")
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)


def chart_pareto(baseline_rows: dict, out: Path) -> None:
    fig, ax = plt.subplots(figsize=(8.5, 5.2))
    colors = {
        "静态本地": "#94a3b8",
        "先到先服务": "#64748b",
        "最小延迟": "#1d4e89",
        "最小成本": "#b45309",
        "遗传算法": "#7c2d12",
        METHOD: "#0f766e",
    }
    for name, rows in baseline_rows.items():
        xs, ys = [], []
        for r in rows:
            if r.get("status") == "SCHEDULED" and r.get("cost") is not None and r.get("latency_ms") is not None:
                xs.append(float(r["cost"]))
                ys.append(float(r["latency_ms"]))
        if not xs:
            continue
        ax.scatter(
            xs,
            ys,
            s=28 if "本文" in name else 18,
            alpha=0.75,
            label=short_name(name),
            c=colors.get(name, "#333"),
        )
    ax.set_xlabel("成本 (元)")
    ax.set_ylabel("时延 (ms)")
    ax.set_title("成功任务成本—时延分布")
    ax.legend(fontsize=8)
    fig.tight_layout()
    fig.savefig(out, dpi=160, bbox_inches="tight")
    plt.close(fig)


def build_report(payload: dict) -> Path:
    setup_font()
    CHARTS.mkdir(parents=True, exist_ok=True)
    summaries = payload["summaries"]
    ablation = payload["ablation"]
    method_rows = payload["paper_rows"]
    method = summaries[METHOD]
    issues = assert_experiment_health(payload)

    paths = {
        "bars": CHARTS / "algo_bars.png",
        "radar": CHARTS / "algo_radar.png",
        "dist": CHARTS / "algo_dist.png",
        "abl": CHARTS / "algo_ablation.png",
        "pareto": CHARTS / "algo_pareto.png",
    }
    chart_bars(summaries, paths["bars"])
    chart_radar(summaries, paths["radar"])
    chart_dist(method_rows, paths["dist"])
    chart_ablation(ablation, paths["abl"])
    chart_pareto(payload.get("baseline_rows") or {}, paths["pareto"])

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.2)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)

    add_para(doc, "跨境算力调度算法测试报告", size=18, bold=True, center=True)
    add_para(doc, "自适应动态权重多目标调度", size=13, center=True)
    add_para(doc, f"版本：最终版　日期：{datetime.now().strftime('%Y-%m-%d')}", size=10, center=True, space_after=14)

    add_heading(doc, "1. 概述", 1)
    add_para(
        doc,
        "本报告给出跨境算力联合调度引擎的实现说明与系统测试结果。"
        "调度目标是在多区域、多约束条件下，为算力任务选择可行且综合更优的落地节点。",
    )
    add_para(
        doc,
        "算法流程为：硬约束过滤（GPU 容量、时延上限、预算、TEE）→ "
        "对可行集做最小—最大规范化 → "
        "按 Score = wl·S(t)·Nlat + wc·Ncost + we·Nenergy + wld·Load 打分选点。"
        "默认权重 wl=0.733，wc=we=wld=0.1，策略整体偏向时延质量。",
    )
    add_para(
        doc,
        "对比基线包括：静态本地、先到先服务（FCFS）、最小延迟、最小成本、遗传算法（GA）。"
        "资源模型覆盖重庆、海南、香港、新加坡、新疆五节点。"
        f"GPU 利用率按运行期峰值占用统计。"
        f"{'实验健康检查通过。' if not issues else '健康检查提示：' + '；'.join(issues)}",
    )

    add_heading(doc, "2. 系统实现", 1)
    add_table(
        doc,
        ["模块", "路径", "说明"],
        [
            ["调度核心", "src/scheduler/adaptive_scheduler.py", "硬约束、规范化、自适应打分"],
            ["实验与基线", "src/scheduler/experiment.py", "30 任务顺序提交、消融与失败分类"],
            ["单元测试", "tests/test_scheduler.py", "约束、打分公式、桥接与实验健康检查"],
            ["控制面", "src/controller/server.py", "API、策略选择、真实派发"],
            ["节点桥接", "src/controller/scheduler_bridge.py", "真实/仿真节点统一为调度状态"],
            ["推理负载", "src/poc_bundle/", "ResNet 分片推理与结果汇聚"],
        ],
    )
    add_para(
        doc,
        "线上控制面默认策略为「动态权重多目标」。"
        "海南、重庆通过 node agent 承接真实 GPU 执行；香港、新加坡、新疆参与混合资源织物下的选点与对照实验。",
    )

    add_heading(doc, "3. 对比实验结果", 1)
    ranking = payload.get("ranking_notes") or {}
    tied_others = ranking.get("success_rate_tied_others") or [
        n for n in (ranking.get("success_rate_tied_with") or []) if not str(n).startswith("本文")
    ]
    add_para(
        doc,
        f"在 30 任务争用场景下，动态权重多目标调度成功率 {method['success_rate_pct']:.2f}%"
        f"（{method['success_tasks']}/{method['total_tasks']}），"
        f"平均时延 {method['avg_latency_ms']:.2f} ms，平均成本 {method['avg_cost']:.2f} 元，"
        f"平均能耗指数 {method['avg_energy']:.2f}，峰值 GPU 利用率 {method['gpu_util_pct']:.1f}%，"
        f"单次决策耗时 {method['avg_compute_ms']:.4f} ms。"
        + (
            f"成功率与{'、'.join(short_name(n) for n in tied_others)}同处最高档；"
            f"同档内本方法时延与成本更优。"
            if tied_others
            else ""
        )
        + (f"未调度任务：{', '.join(method['failed_ids'])}。" if method.get("failed_ids") else "全部任务成功调度。"),
    )
    rows = []
    for name, s in summaries.items():
        rows.append([
            short_name(name),
            s["success_tasks"],
            s["unscheduled_tasks"],
            f"{s['success_rate_pct']:.2f}",
            f"{s['avg_latency_ms']:.2f}",
            f"{s['avg_cost']:.2f}",
            f"{s['avg_energy']:.2f}",
            f"{s['gpu_util_pct']:.1f}",
            f"{s['avg_compute_ms']:.4f}",
        ])
    add_table(
        doc,
        ["算法", "成功", "未调度", "成功率%", "平均时延", "平均成本", "平均能耗", "峰值利用率%", "计算ms"],
        rows,
    )
    add_figure(doc, paths["bars"], "图1  关键指标对比")
    add_figure(doc, paths["radar"], "图2  多维度雷达对比")
    add_figure(doc, paths["dist"], "图3  动态权重多目标五节点落点分布")
    add_figure(doc, paths["pareto"], "图4  成功任务成本—时延分布")

    diffs = payload.get("decision_diff_vs_paper") or {}
    if diffs:
        add_para(doc, "相对本方法的决策差异（选点或成败不同计为差异）：")
        add_table(
            doc,
            ["对比算法", "相同决策数", "差异率%"],
            [[short_name(k), v["identical_decisions"], v["diff_rate_pct"]] for k, v in diffs.items()],
        )

    ga = summaries.get("遗传算法", {})
    if ga and method["avg_compute_ms"] > 0:
        speedup = ga["avg_compute_ms"] / max(method["avg_compute_ms"], 1e-9)
        add_para(
            doc,
            f"相对遗传算法，在线打分平均加速约 {speedup:.0f} 倍"
            f"（GA {ga['avg_compute_ms']:.2f} ms，本方法 {method['avg_compute_ms']:.4f} ms）。",
        )

    add_heading(doc, "3.1 失败原因分类", 1)
    tax = payload.get("failure_taxonomy") or {}
    counts = tax.get("counts") or {}
    by_task = tax.get("by_task") or {}
    if counts:
        add_table(doc, ["失败类型", "任务数"], [[k, v] for k, v in counts.items()])
    if by_task:
        add_table(doc, ["任务", "归类"], [[tid, reason] for tid, reason in by_task.items()])
    cold = tax.get("cold_start") or {}
    if cold:
        add_para(
            doc,
            "空集群即不可行的任务（与运行争用无关）："
            + "；".join(f"{k}（{v}）" for k, v in cold.items())
            + "。其余失败主要来自顺序提交下的容量争用。",
        )

    add_heading(doc, "4. 消融实验", 1)
    add_para(doc, "关闭单项权重后，成功率、时延、成本及失败集变化如下。")
    ab_rows = []
    for name, s in ablation.items():
        ab_rows.append([
            short_name(name),
            f"{s['success_rate_pct']:.2f}",
            f"{s['avg_latency_ms']:.2f}",
            f"{s['avg_cost']:.2f}",
            f"{s['avg_energy']:.2f}",
            ", ".join(s.get("failed_ids") or []) or "—",
        ])
    add_table(doc, ["配置", "成功率%", "平均时延", "平均成本", "平均能耗", "失败任务"], ab_rows)
    add_figure(doc, paths["abl"], "图5  消融实验指标")
    for note in payload.get("ablation_notes") or []:
        add_para(doc, f"• {note}", size=10)

    add_heading(doc, "5. 真实联调说明", 1)
    add_para(
        doc,
        "控制面默认使用本算法为海南/重庆真实节点选点，并保留 16GB 显存等硬约束桥接逻辑。"
        "现场 ResNet 分片推理、短时占卡观测利用率、结果汇聚与安全负向用例见"
        "《跨境算力调度真实联调测试报告》。",
    )

    add_heading(doc, "6. 结论", 1)
    static = summaries.get("静态本地", {})
    mincost = summaries.get("最小成本", {})
    add_para(
        doc,
        f"测试表明，自适应动态权重多目标调度在本场景成功率达到 {method['success_rate_pct']:.2f}%，"
        f"优于静态本地（{static.get('success_rate_pct', '—')}%）"
        + (f"与最小成本（{mincost.get('success_rate_pct')}%）" if mincost else "")
        + f"；在最高成功率档内取得更低时延与成本"
        f"（{method['avg_latency_ms']:.2f} ms / {method['avg_cost']:.2f} 元），"
        f"峰值利用率 {method['gpu_util_pct']:.1f}%，决策时延 {method['avg_compute_ms']:.4f} ms，"
        "满足跨境算力联合调度的在线决策需求。",
    )

    out = REPORTS / "跨境算力调度算法测试报告.docx"
    doc.save(out)
    (DATA / "experiment_meta.json").write_text(
        json.dumps(
            {
                "generated_at": datetime.now().isoformat(timespec="seconds"),
                "report": str(out.name),
                "summaries": summaries,
                "ablation": ablation,
                "ablation_notes": payload.get("ablation_notes"),
                "failure_taxonomy": payload.get("failure_taxonomy"),
                "ranking_notes": payload.get("ranking_notes"),
                "meta": payload.get("meta"),
                "health_issues": issues,
                "decision_diff": diffs,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    return out


def build_integration_report() -> Path:
    """生成与当前线上系统一致的真实联调说明报告（不含凭据）。"""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    add_para(doc, "跨境算力调度真实联调测试报告", size=18, bold=True, center=True, space_after=4)
    add_para(
        doc,
        f"版本：最终版　日期：{datetime.now().strftime('%Y-%m-%d')}　织物：real（海南 / 重庆）",
        size=10,
        center=True,
        space_after=16,
    )
    add_para(
        doc,
        "本报告描述新加坡控制面与两地真实 GPU agent 的联调现状。算法为自适应动态权重多目标调度；"
        "控制台提供利用率峰值、收益对照与场景化测试。凭据与内网细节不写入本报告。",
    )

    add_heading(doc, "1. 联调拓扑", 1)
    add_table(
        doc,
        ["角色", "位置", "职责"],
        [
            ["控制面 / UI", "新加坡 :8080", "场景提交、调度决策、状态汇聚、收益面板"],
            ["Node Agent", "海南", "nvidia-smi 上报、ResNet / gpu_load 执行"],
            ["Node Agent", "重庆", "同上，多卡并行"],
        ],
    )

    add_heading(doc, "2. 调度与负载", 1)
    add_para(
        doc,
        "默认策略「动态权重多目标」。DT01 等场景先跑 ResNet 分片（通常数秒），再按 duration_sec "
        "短时 GEMM 占卡（最多约 15 秒）便于观察利用率；DEMO 场景使用持续 gpu_load≈25 秒。"
        "控制面记录区域 / 单卡会话峰值利用率，任务结束后瞬时利用率归零属预期行为。",
    )

    add_heading(doc, "3. 主要验收点", 1)
    add_table(
        doc,
        ["编号", "场景", "关注点"],
        [
            ["DT01", "两地分片批量推理", "跨区落点、结果汇聚、利用率峰值"],
            ["DT02", "海南优先与分流", "显存占满后改派重庆"],
            ["DT03", "16GB 显存", "排除重庆 12GB 卡"],
            ["DEMO", "利用率与收益演示", "持续占卡、成本/时延收益摘要"],
            ["TC02/03/04", "安全负向", "鉴权/端口/哈希拒绝且不落 GPU"],
        ],
    )

    add_heading(doc, "4. 控制台能力", 1)
    add_para(
        doc,
        "实时折线：可用显存 / GPU 利用率（含会话峰）/ 分片队列；资源卡片展示实时与峰值利用率、显存占用；"
        "完成后「测试收益与效果」汇总相对单边基线的成本与时延收益。",
    )

    add_heading(doc, "5. 结论", 1)
    add_para(
        doc,
        "真实联调路径已收敛为：新算法 + 经典控制台 + 两地真实 agent。"
        "算法报告给出离线 30 任务对比；本报告对应现场可复现的调度与观测能力。",
    )

    out = REPORTS / "跨境算力调度真实联调测试报告.docx"
    doc.save(out)
    return out


def main() -> None:
    DATA.mkdir(parents=True, exist_ok=True)
    CHARTS.mkdir(parents=True, exist_ok=True)
    payload = run_paper_experiment(DATA)
    algo = build_report(payload)
    integ = build_integration_report()
    print(algo)
    print(integ)
    issues = assert_experiment_health(payload)
    print("health:", "OK" if not issues else issues)


if __name__ == "__main__":
    main()
