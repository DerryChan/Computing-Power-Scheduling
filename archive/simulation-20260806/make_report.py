"""Generate the formal Chinese test report from outputs produced by l1_scheduler_test.py."""

from __future__ import annotations

import json
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "outputs"
CHARTS = OUT / "charts"
DOCX_PATH = ROOT / "L1_算力调度仿真测试报告_20260806.docx"
MD_PATH = ROOT / "L1_算力调度仿真测试报告_20260806.md"
ZIP_PATH = ROOT / "L1_算力调度测试证据包_20260806.zip"

BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "F3F8FC"
GRAY = "666666"
GREEN = "548235"
ORANGE = "BF9000"
RED = "C00000"
FONT_NAME = "Noto Sans SC Thin"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str = "000000", size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(cell, top=100, start=100, bottom=100, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table, color="D9E2F3", size="6") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None, font_size: int = 8) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    set_row_cant_split(hdr)
    for i, header in enumerate(headers):
        set_cell_text(hdr.cells[i], header, bold=True, color="FFFFFF", size=font_size)
        shade(hdr.cells[i], BLUE)
        set_cell_margins(hdr.cells[i])
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
            if ridx % 2 == 0:
                shade(cells[i], PALE_BLUE)
            set_cell_margins(cells[i])
        if widths:
            for i, width in enumerate(widths):
                cells[i].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.color.rgb = RGBColor.from_string(BLUE if level == 1 else "365F91")
    return p


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.color.rgb = RGBColor.from_string(BLUE)
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    p.add_run(text)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t, color=fill, size="8")
    set_repeat_table_header(t.rows[0])
    cell = t.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, top=160, start=180, bottom=160, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_figure(doc: Document, path: Path, caption: str, width_cm: float = 16.5) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inline = p.add_run().add_picture(str(path), width=Cm(width_cm))
    # Add a meaningful alternative description for accessibility and document QA.
    inline._inline.docPr.set("descr", caption)
    inline._inline.docPr.set("title", caption)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(6)
    r = cp.add_run(caption)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(GRAY)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    paragraph.add_run(" 页")


def configure_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.1)
    sec.right_margin = Cm(2.1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("222222")
    for name, size in [("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 10.5)]:
        st = styles[name]
        st.font.name = FONT_NAME
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(BLUE)
    footer = sec.footer.paragraphs[0]
    footer.text = "L1 跨境算力调度仿真测试报告  |  2026-08-06  |  "
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    add_page_number(sec.footer.add_paragraph())


def make_docx(data: dict, cases: pd.DataFrame, strategy: pd.DataFrame, snapshots: pd.DataFrame) -> None:
    doc = Document()
    configure_doc(doc)
    pass_count = int((cases["status"] == "PASS-SIM").sum())
    not_run_count = int((cases["status"] == "NOT-RUN").sum())
    not_applicable_count = int((cases["status"] == "NOT-APPLICABLE").sum())
    unresolved_a = cases[cases["level"].astype(str).str.contains("A") & (cases["status"] != "PASS-SIM")]

    # Cover page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(100)
    r = p.add_run("跨境算力调度 PoC")
    r.font.size = Pt(17)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    r = p2.add_run("L1 算力调度仿真测试报告")
    r.font.size = Pt(25)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("17365D")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run("依据《跨境算力调度 PoC 测试实施方案 v0724》编制")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(35)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["报告编号", "L1-SCHED-SIM-20260806"],
            ["执行日期", "2026年8月6日"],
            ["测试性质", "本地确定性仿真；不连接真实服务器、VPN、OTN或调度控制面"],
            ["资源边界", "海南 RTX 4090 24GB×2；重庆 RTX 4070 12GB×4；海南—重庆 OTN按100 Mbit/s建模"],
            ["测试依据", "测试实施方案 v0724、重庆测试服务器资源说明、GitHub参考调度脚本"],
        ],
        widths=[3.0, 13.0],
        font_size=9,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("内部测试材料 · 凭据和受控网络信息未写入本报告")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(GRAY)

    add_heading(doc, "执行摘要", 1)
    add_callout(doc, "仿真结论", f"用例矩阵已覆盖方案中的TC01-TC30和DT01-DT06：{pass_count}项PASS-SIM，{not_run_count}项真实现场待执行，{not_applicable_count}项因条件未满足暂不适用。仿真未发现调度规则、单卡显存约束、两地分片、异常停派、结果隔离和清理逻辑方面的失败；由于仍有{len(unresolved_a)}项A类用例需要真实环境验证，本报告不形成现场验收通过结论。", fill="E2F0D9")
    add_body(doc, "本次测试以方案中的L1多地轻量验证环境为边界，复现“新加坡任务发起—海南接入与调度—海南/重庆独立执行—海南汇聚—结果回传与清理”的任务级并行方式。测试数据为4096条确定性合成样本，划分为8个分片，每个子任务占用单卡、8GB显存。")
    add_table(
        doc,
        ["结果项", "本次结果", "与方案要求的关系"],
        [
            ["用例覆盖", f"{len(cases)}项（TC01-TC30、DT01-DT06）", "已建立逐项状态和证据记录"],
            ["仿真用例", f"{pass_count}通过 / {(cases['status'] == 'FAIL-SIM').sum()}失败", "PASS-SIM仅表示本地逻辑仿真通过"],
            ["现场待执行", f"{not_run_count}项", "需要真实认证、网络、节点代理或调度控制面"],
            ["条件未满足", f"{not_applicable_count}项", "TC16/DT06需训练镜像、数据和许可冻结后执行"],
            ["DT01主流程", "3轮×8分片；每轮4096条；两地均参与", "满足方案核心业务任务的逻辑要求"],
            ["DT02/DT03调度约束", "海南优先；显存不足分流；16GB任务排除4070", "满足方案调度规则验证要求"],
            ["DT04异常恢复", "断链后重庆停派；失败分片重试至海南", "满足方案异常处理逻辑要求"],
            ["TC30稳定性", "192项任务，24小时虚拟时间，100%成功", "仅为加速仿真，不等价于现场24小时运行"],
        ],
        widths=[3.0, 7.0, 6.0],
        font_size=8.5,
    )

    add_heading(doc, "1. 测试依据与范围", 1)
    add_body(doc, "测试方案明确：本期仅做L1验证，不将海南4090与重庆4070拼接为统一显存池，不开展跨站点模型并行、张量并行或梯度同步；两地共同计算采用“父任务拆分—两地独立执行—海南汇总”的任务级并行方式。")
    add_body(doc, "调度参考实现采用用户指定的GitHub脚本 scheduler_simulation_improved.py，并在其多策略比较、消融指标和可视化结构基础上，补充了L1资源边界、父子任务、分片、单卡显存、链路状态、重试、隔离、清理和审计字段。参考版本：b18138d88098f4a5db31f09c81e57e451fd79377。")
    add_body(doc, "参考链接：https://github.com/derry-cheng/test/blob/main/scheduler_simulation_improved.py")

    add_heading(doc, "2. 仿真环境与建模边界", 1)
    add_table(
        doc,
        ["区域/组件", "本次建模配置", "状态"],
        [
            ["新加坡任务端", "任务提交、输入哈希登记、结果接收逻辑抽象", "仿真"],
            ["海南", "RTX 4090 24GB×2；作为接入、调度、本地执行和汇聚节点", "仿真"],
            ["重庆", "RTX 4070 12GB×4；作为异地执行节点", "仿真"],
            ["海南—重庆", "OTN 100 Mbit/s；用于分片传输和链路中断建模", "仿真参数"],
            ["新加坡—海南", "VPN上传/回传流程抽象；吞吐、丢包和抖动未测", "现场待执行"],
        ],
        widths=[3.2, 8.5, 4.3],
        font_size=8.5,
    )
    add_figure(doc, CHARTS / "resource_snapshot.png", "图1  单卡显存资源快照与DT03的16GB硬约束", 15.5)
    add_body(doc, "图1显示：海南每张4090可提供24GB显存，重庆每张4070为12GB。对单卡16GB任务，重庆在候选节点阶段被直接排除；不能把4张4070的总显存视为一张48GB显存。")

    add_heading(doc, "3. 测试任务与执行方法", 1)
    add_table(
        doc,
        ["项目", "冻结/建模内容"],
        [
            ["父任务", "POC-IMG-DUAL-001至003；DT01连续执行3轮"],
            ["样本与分片", "4096条合成样本；8个分片；每片512条；sample_id唯一"],
            ["子任务资源", "GPU×1；显存需求8GB；CPU 4核、内存8GB按逻辑字段保留"],
            ["节点策略", "满足硬约束后海南优先；支持海南/重庆亲和；重庆链路不可用时从候选集剔除"],
            ["结果校验", "分片结果哈希、样本数、唯一性、缺失数、父子任务绑定和A/B任务隔离"],
            ["稳定性模型", "24个虚拟小时、每15分钟2项任务，共192项；不替代真实24小时测试"],
        ],
        widths=[3.2, 13.0],
        font_size=8.5,
    )
    add_figure(doc, CHARTS / "dispatch_timeline.png", "图2  DT01首轮8个子任务的跨区域执行时间线", 16.5)

    add_heading(doc, "4. 测试结果", 1)
    add_heading(doc, "4.1 方案核心场景DT01—DT05", 2)
    dt_rows = [
        ["DT01", "两地共同执行的分片批量推理", "PASS-SIM", "连续3轮；每轮8分片/4096样本；海南和重庆均执行；缺失0、重复0"],
        ["DT02", "海南优先与重庆自动分流", "PASS-SIM", "空闲任务→海南；海南每卡预占20GB后8GB任务→重庆"],
        ["DT03", "单卡16GB显存硬约束", "PASS-SIM", "重庆12GB单卡候选阶段排除；海南离线时不向重庆错误派发"],
        ["DT04", "重庆链路中断、停派与分片重试", "PASS-SIM", "运行中重庆分片失败；未开始分片停止派往重庆；重试至海南"],
        ["DT05", "结果汇聚、跨境回传与数据清理", "PASS-SIM", "A/B任务独立结果键和哈希；模拟确认后清理4类非审计对象"],
    ]
    add_table(doc, ["编号", "场景", "判定", "实测/仿真证据"], dt_rows, widths=[1.6, 5.0, 2.2, 7.8], font_size=8.2)
    add_figure(doc, CHARTS / "dt02_routing.png", "图3  DT02海南优先与显存不足后的重庆分流", 13.5)
    add_figure(doc, CHARTS / "failure_recovery.png", "图4  DT04链路中断、失败、停派与重试状态", 16.5)

    add_heading(doc, "4.2 调度策略对比", 2)
    strategy_labels = {"静态本地": "静态本地", "先到先服务": "先到先服务", "海南优先": "海南优先", "最小延迟": "最小延迟", "最小成本": "最小成本", "加权平均": "加权平均"}
    strategy_rows = []
    for _, r in strategy.iterrows():
        strategy_rows.append([
            strategy_labels.get(r["algorithm"], r["algorithm"]),
            f"{r['success_rate_pct']:.2f}%",
            f"{r['avg_latency_ms']:.2f}",
            f"{r['avg_cost']:.2f}",
            f"{r['unscheduled_tasks']:.0f}",
        ])
    add_table(doc, ["策略", "成功率", "平均延迟(ms)", "平均成本", "未调度任务"], strategy_rows, widths=[4.0, 2.4, 3.2, 2.6, 3.0], font_size=8.2)
    add_body(doc, "对60项合成工作负载的结果显示：静态本地策略有7项16GB任务被错误固定到重庆后无法调度，成功率为88.33%；动态策略均达到100%。最小成本策略的平均成本最低，但平均延迟最高，体现跨区域调度中的成本—时延权衡。上述成本、时延和能耗均为模型参数，不代表现场计费或网络实测值。")
    add_figure(doc, CHARTS / "strategy_comparison.png", "图5  合成工作负载下的多策略对比", 16.5)

    add_heading(doc, "4.3 稳定性与审计", 2)
    add_table(
        doc,
        ["指标", "结果", "说明"],
        [
            ["TC30虚拟运行时长", "24小时", "按15分钟时间槽加速仿真"],
            ["混合任务数", "192项", "8GB与16GB任务混合；海南优先/加权平均混合"],
            ["成功率", "100.00%", "高于方案压力任务95%门槛，仅为仿真结果"],
            ["状态丢失", "0", "每项任务均有状态记录"],
            ["数据错配", "0", "父任务/子任务/结果键隔离"],
            ["审计必填字段", "100%", "task_id、stage、timestamp、region、status、reason"],
        ],
        widths=[4.0, 4.0, 8.0],
        font_size=8.5,
    )
    add_figure(doc, CHARTS / "stability.png", "图6  TC30 24小时加速稳定性仿真", 16.5)

    add_heading(doc, "5. 用例执行矩阵", 1)
    executed = cases[cases["status"] == "PASS-SIM"]
    executed_rows = [[r["case_id"], r["level"], r["title"], r["status"], r["coverage_source"], r["metrics"]] for _, r in executed.iterrows()]
    add_table(doc, ["用例", "级别", "场景", "判定", "覆盖来源", "结果摘要"], executed_rows, widths=[1.5, 1.2, 4.0, 1.6, 2.4, 5.1], font_size=7.4)
    add_heading(doc, "5.1 当前未执行用例", 2)
    not_run = cases[cases["status"].isin(["NOT-RUN", "NOT-APPLICABLE"])]
    not_rows = [[r["case_id"], r["level"], r["title"], r["status"], r["detail"]] for _, r in not_run.iterrows()]
    add_table(doc, ["用例", "级别", "场景", "状态", "原因"], not_rows, widths=[1.5, 1.2, 4.0, 2.0, 7.0], font_size=7.6)
    add_callout(doc, "现场验收边界", f"当前仍有{len(unresolved_a)}项A类用例未在真实环境执行：{', '.join(unresolved_a['case_id'].tolist())}。按照方案“A类用例应100%通过”的验收口径，在这些项目完成前，整体结论只能记为“仿真验证通过，现场验收待执行”。", fill="FFF2CC")

    add_heading(doc, "6. 证据文件与可复测性", 1)
    add_table(
        doc,
        ["证据类型", "文件", "用途"],
        [
            ["测试脚本", "l1_scheduler_test.py", "重新生成全部测试数据和图表"],
            ["用例矩阵", "outputs/case_results.csv", "用例级别、状态、证据和结果摘要"],
            ["逐项证据", "outputs/case_evidence/TC01.json至DT06.json", "覆盖方案要求的父子任务、输入哈希、候选集、选择原因、输出、清理、结论等字段"],
            ["覆盖矩阵", "outputs/coverage_matrix.csv", "TC01-TC30、DT01-DT06逐项映射及现场待填字段"],
            ["DT01", "outputs/dt01_shards.csv; outputs/report_data.json", "三轮父子任务、分片和样本完整性"],
            ["DT02/DT03", "outputs/dt02_routing.csv; outputs/dt03_vram_constraint.csv", "海南优先、分流和单卡显存约束"],
            ["DT04/DT05", "outputs/dt04_failure_recovery.csv; outputs/dt05_isolation_cleanup.csv", "故障恢复、隔离、结果和清理"],
            ["稳定性", "outputs/tc30_stability.csv", "24小时虚拟时间槽任务记录"],
            ["可视化", "outputs/charts/*.png", "策略、时间线、显存、故障、稳定性图表"],
        ],
        widths=[3.0, 7.0, 6.0],
        font_size=8.2,
    )
    add_body(doc, f"本次合成输入包冻结哈希：{data['input_sha256']}。报告和证据包未包含服务器密码、私钥、VPN参数或受控网络地址。")

    add_heading(doc, "7. 现场复测建议", 1)
    for text in [
        "在海南、重庆节点代理和调度控制面上线后，首先执行资源注册、GPU/DCGM核对、镜像启动和单节点推理预验证。",
        "冻结真实镜像摘要、模型权重、约140MB合成数据包、manifest.csv、8个分片清单和参考输出；以本报告中的父子任务字段作为现场记录模板。",
        "补执行TC02、TC03、TC24，并按方案采集VPN/OTN吞吐、时延、丢包、抖动、真实状态刷新和审计截图；TC16/DT06在训练条件满足后单独执行。",
        "在真实环境连续执行DT01三轮，确认海南、重庆均实际使用GPU；再执行DT04受控链路中断和DT05双父任务隔离清理。",
        "现场24小时稳定性测试必须以真实墙钟时间、真实任务日志、GPU监控和链路监控为准；本报告的TC30加速仿真只能作为执行前的逻辑回归。",
        "将现场测试结果回填到逐项用例矩阵和JSON证据；只有A类全部通过、核心流程连续3次成功、清理和证据归档完整后，才可依据方案形成正式验收结论。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "8. 结论", 1)
    add_callout(doc, "结论：仿真验证通过；现场验收待执行", "在限定的L1资源模型和合成任务下，调度逻辑能够实现两地分片执行、海南优先、重庆分流、单卡显存硬约束、链路中断停派与重试、结果隔离、清理和审计字段生成。由于本次没有连接真实服务器、VPN、OTN、调度API和节点代理，不能据此证明实际系统的网络性能、GPU执行正确性、身份访问控制或生产级稳定性。", fill="E2F0D9")
    add_body(doc, "本报告结论只适用于本次仿真的资源边界、任务规模和合成数据，不外推为生产级服务能力或商业SLA。")
    doc.save(DOCX_PATH)


def make_markdown(data: dict, cases: pd.DataFrame, strategy: pd.DataFrame) -> None:
    pass_count = int((cases["status"] == "PASS-SIM").sum())
    not_run_count = int((cases["status"] == "NOT-RUN").sum())
    not_applicable_count = int((cases["status"] == "NOT-APPLICABLE").sum())
    unresolved_a = cases[cases["level"].astype(str).str.contains("A") & (cases["status"] != "PASS-SIM")]
    lines = []
    lines += [
        "# 跨境算力调度 PoC：L1 算力调度仿真测试报告",
        "",
        "> 报告编号：L1-SCHED-SIM-20260806  \\",
        "> 执行日期：2026-08-06  \\",
        "> 测试性质：本地确定性仿真，不连接真实服务器、VPN、OTN或调度控制面",
        "",
        "## 结论",
        "",
        f"**仿真验证通过；现场验收待执行。** 用例矩阵已覆盖TC01-TC30和DT01-DT06：{pass_count}项PASS-SIM，{not_run_count}项真实现场待执行，{not_applicable_count}项条件未满足；当前仍有{len(unresolved_a)}项A类用例需要真实环境验证。",
        "",
        "| 项目 | 结果 |",
        "|---|---|",
        f"| 用例覆盖 | {len(cases)}项（TC01-TC30、DT01-DT06） |",
        f"| 仿真用例 | {pass_count}通过 / {int((cases['status'] == 'FAIL-SIM').sum())}失败 |",
        f"| 现场待执行 | {not_run_count}项 |",
        f"| 条件未满足 | {not_applicable_count}项 |",
        "| DT01 | 3轮×8分片；每轮4096条；两地均参与 |",
        "| DT02 | 海南优先；海南显存不足时分流重庆 |",
        "| DT03 | 16GB任务排除重庆12GB单卡 |",
        "| DT04 | OTN中断后重庆停派，失败分片重试海南 |",
        "| DT05 | A/B父任务隔离、结果哈希和清理逻辑通过 |",
        "| TC30 | 192项任务；24小时虚拟时间；100%成功 |",
        "",
        "## 1. 范围与边界",
        "",
        "测试依据为《跨境算力调度 PoC 测试实施方案 v0724》、重庆测试服务器资源说明和 GitHub 参考脚本。资源边界为海南 RTX 4090 24GB×2、重庆 RTX 4070 12GB×4，海南—重庆 OTN按100 Mbit/s建模。",
        "",
        "本次未测量真实VPN/OTN吞吐、时延、丢包、抖动、GPU利用率、真实24小时稳定性、认证控制和真实调度接口。",
        "",
        "参考脚本：[scheduler_simulation_improved.py](https://github.com/derry-cheng/test/blob/main/scheduler_simulation_improved.py)，参考版本 `b18138d88098f4a5db31f09c81e57e451fd79377`。",
        "",
        "## 2. 核心场景结果",
        "",
        "| 编号 | 场景 | 判定 | 结果摘要 |",
        "|---|---|---|---|",
        "| DT01 | 两地共同执行的分片批量推理 | PASS-SIM | 连续3轮；每轮8分片/4096样本；缺失0、重复0 |",
        "| DT02 | 海南优先与重庆自动分流 | PASS-SIM | 空闲→海南；海南每卡预占20GB后8GB任务→重庆 |",
        "| DT03 | 单卡16GB显存硬约束 | PASS-SIM | 重庆12GB候选阶段排除；海南离线时不错误派发 |",
        "| DT04 | 重庆链路中断、停派与分片重试 | PASS-SIM | 重庆运行分片失败；未开始分片停派；重试至海南 |",
        "| DT05 | 结果汇聚、回传与清理 | PASS-SIM | A/B结果键隔离；哈希分离；清理4类对象 |",
        "",
        "![策略对比](outputs/charts/strategy_comparison.png)",
        "",
        "![DT01时间线](outputs/charts/dispatch_timeline.png)",
        "",
        "![DT03显存约束](outputs/charts/resource_snapshot.png)",
        "",
        "![DT04故障恢复](outputs/charts/failure_recovery.png)",
        "",
        "![TC30稳定性](outputs/charts/stability.png)",
        "",
        "## 3. 策略对比",
        "",
        "| 策略 | 成功率 | 平均延迟(ms) | 平均成本 | 未调度任务 |",
        "|---|---:|---:|---:|---:|",
    ]
    for _, r in strategy.iterrows():
        lines.append(f"| {r['algorithm']} | {r['success_rate_pct']:.2f}% | {r['avg_latency_ms']:.2f} | {r['avg_cost']:.2f} | {int(r['unscheduled_tasks'])} |")
    lines += [
        "",
        "静态本地策略有7项16GB任务被固定到重庆后无法调度，成功率88.33%；动态策略均达到100%。最小成本策略平均成本最低，但平均延迟最高。成本、延迟和能耗均为建模指标。",
        "",
        "## 4. 用例状态",
        "",
        "| 用例 | 级别 | 场景 | 状态 | 覆盖来源 |",
        "|---|---|---|---|---|",
    ]
    for _, r in cases.iterrows():
        lines.append(f"| {r['case_id']} | {r['level']} | {r['title']} | {r['status']} | {r['coverage_source']} |")
    lines += [
        "",
        f"A类现场待执行项目：{', '.join(unresolved_a['case_id'].tolist())}。TC16和DT06为条件适用项，需在轻量训练条件冻结后单独执行。",
        "",
        "## 5. 逐项证据与现场复测",
        "",
        "每个TC/DT均生成 `outputs/case_evidence/<case_id>.json`，并在 `outputs/coverage_matrix.csv` 中记录父任务ID、子任务ID、执行时间、前置条件、输入文件及SHA-256、分片、候选集、选择原因、节点/GPU、输出、缺失/重复、日志证据、问题编号、重试关系、清理记录和结论。PASS-SIM字段明确标注为仿真，现场待执行项保留“现场待填”。",
        "",
        "1. 冻结真实镜像摘要、模型、合成数据、manifest和参考输出。",
        "2. 完成海南/重庆节点代理、调度API、VPN和OTN部署后，先进行资源注册与单节点预验证。",
        f"3. 补执行{', '.join(unresolved_a['case_id'].tolist())}；DT01现场连续3轮；DT04受控断链；DT05双父任务隔离清理。",
        "4. 真实24小时稳定性必须依据墙钟时间、GPU监控、链路监控和平台日志判定。",
        "5. A类用例全部通过、核心流程连续3次成功、清理和证据归档完整后，再形成正式验收结论。",
        "",
        f"输入包冻结哈希：`{data['input_sha256']}`。凭据和受控网络信息未写入报告。",
    ]
    MD_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")


def make_zip() -> None:
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for path in [DOCX_PATH, MD_PATH, ROOT / "l1_scheduler_test.py", ROOT / "make_report.py"]:
            z.write(path, arcname=path.name)
        for path in sorted(OUT.rglob("*")):
            if path.is_file():
                z.write(path, arcname=str(Path("outputs") / path.relative_to(OUT)))


def main() -> None:
    data = json.loads((OUT / "report_data.json").read_text(encoding="utf-8"))
    cases = pd.read_csv(OUT / "case_results.csv", encoding="utf-8-sig")
    strategy = pd.read_csv(OUT / "strategy_metrics.csv", encoding="utf-8-sig")
    snapshots = pd.read_csv(OUT / "resource_snapshot.csv", encoding="utf-8-sig")
    make_docx(data, cases, strategy, snapshots)
    make_markdown(data, cases, strategy)
    make_zip()
    print(f"created: {DOCX_PATH}")
    print(f"created: {MD_PATH}")
    print(f"created: {ZIP_PATH}")


if __name__ == "__main__":
    main()
